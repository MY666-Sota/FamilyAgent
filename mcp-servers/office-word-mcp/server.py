"""
office-word-mcp — 端口 9001，SSE 传输
用本地 python-docx 直接读写 Word 文档，无需外部 9011 服务。
文件落 shared/outputs/，URL 通过 file-server (8090) 访问。
"""
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from mcp.server.fastmcp import FastMCP

OUTPUTS_DIR = Path(__file__).parent.parent.parent / "shared" / "outputs"
OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)

FILE_SERVER_BASE = os.getenv("FILE_SERVER_BASE", "http://localhost:8090/files")

mcp = FastMCP("office-word", port=9001)


def _file_url(filename: str) -> str:
    return f"{FILE_SERVER_BASE}/{filename}"


def _safe_filename(filename: str) -> str:
    if not filename.endswith(".docx"):
        filename += ".docx"
    # 只允许文件名，不允许路径分隔符
    return Path(filename).name


def _apply_markdown_to_doc(doc: Document, text: str) -> None:
    """把简单 Markdown/纯文本写入 Document。支持 # 标题、**加粗**、- 列表。"""
    for line in text.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "":
            doc.add_paragraph("")
        else:
            # 处理行内 **加粗**
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)


@mcp.tool()
def create_document(
    filename: str,
    content: str,
    title: str = "",
    author: str = "FamilyAgent",
) -> dict:
    """
    创建 Word 文档（.docx）并保存到 shared/outputs/。

    Args:
        filename: 输出文件名，例如 report.docx
        content:  Markdown 或纯文本正文
        title:    文档标题（可选，写入文档属性和首行 H1）
        author:   作者名（可选，写入文档属性）

    Returns:
        {"file_url": str, "filename": str}
    """
    filename = _safe_filename(filename)
    dest = OUTPUTS_DIR / filename

    doc = Document()
    doc.core_properties.author = author
    if title:
        doc.core_properties.title = title
        doc.add_heading(title, level=1)

    _apply_markdown_to_doc(doc, content)
    doc.save(dest)

    return {"file_url": _file_url(filename), "filename": filename}


@mcp.tool()
def read_document(filename: str) -> dict:
    """
    读取 shared/outputs/ 中的 Word 文档，返回纯文本内容。

    Args:
        filename: 文件名，例如 report.docx

    Returns:
        {"content": str, "filename": str}
    """
    filename = _safe_filename(filename)
    target = OUTPUTS_DIR / filename
    if not target.exists():
        raise FileNotFoundError(f"文件不存在: {filename}")

    doc = Document(str(target))
    lines = [para.text for para in doc.paragraphs]
    return {"content": "\n".join(lines), "filename": filename}


@mcp.tool()
def list_documents() -> dict:
    """
    列出 shared/outputs/ 中所有 .docx 文件。

    Returns:
        {"files": [{"filename": str, "file_url": str}]}
    """
    files = [
        {"filename": f.name, "file_url": _file_url(f.name)}
        for f in OUTPUTS_DIR.glob("*.docx")
    ]
    return {"files": files}


@mcp.tool()
def append_to_document(filename: str, content: str) -> dict:
    """
    向已有 Word 文档追加内容。

    Args:
        filename: 目标文件名
        content:  要追加的 Markdown 或纯文本

    Returns:
        {"file_url": str, "filename": str}
    """
    filename = _safe_filename(filename)
    target = OUTPUTS_DIR / filename
    if not target.exists():
        raise FileNotFoundError(f"文件不存在: {filename}")

    doc = Document(str(target))
    _apply_markdown_to_doc(doc, content)
    doc.save(target)

    return {"file_url": _file_url(filename), "filename": filename}


if __name__ == "__main__":
    mcp.run(transport="sse")
